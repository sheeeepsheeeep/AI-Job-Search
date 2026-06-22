"""
CV / Resume text extraction service.
Supports PDF (via PyPDF2) and DOCX (via python-docx) formats.
"""

import logging
import os
from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text content from a PDF file.

    Args:
        file_path: Absolute or relative path to the PDF file.

    Returns:
        Concatenated text from all pages.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file cannot be parsed as a valid PDF.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")

    try:
        reader = PdfReader(str(path))
        pages_text: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text.strip())
        full_text = "\n\n".join(pages_text)
        logger.info("Extracted %d characters from PDF (%d pages): %s", len(full_text), len(reader.pages), file_path)
        return full_text
    except Exception as exc:
        logger.error("Failed to extract text from PDF %s: %s", file_path, exc)
        raise ValueError(f"Could not parse PDF file: {exc}") from exc


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text content from a DOCX file.

    Args:
        file_path: Absolute or relative path to the DOCX file.

    Returns:
        Concatenated text from all paragraphs.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file cannot be parsed as a valid DOCX.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    try:
        doc = Document(str(path))
        paragraphs_text: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs_text.append(text)

        # Also extract text from tables (CVs sometimes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_texts: list[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    paragraphs_text.append(" | ".join(row_texts))

        full_text = "\n".join(paragraphs_text)
        logger.info("Extracted %d characters from DOCX: %s", len(full_text), file_path)
        return full_text
    except Exception as exc:
        logger.error("Failed to extract text from DOCX %s: %s", file_path, exc)
        raise ValueError(f"Could not parse DOCX file: {exc}") from exc


def extract_cv_text(file_path: str) -> str:
    """Dispatcher – extract text from a CV file based on its extension.

    Supported formats: .pdf, .docx

    Args:
        file_path: Path to the CV file.

    Returns:
        Extracted text content.

    Raises:
        ValueError: If the file extension is not supported.
    """
    ext = os.path.splitext(file_path)[1].lower()

    extractors = {
        ".pdf": extract_text_from_pdf,
        ".docx": extract_text_from_docx,
    }

    extractor = extractors.get(ext)
    if extractor is None:
        supported = ", ".join(extractors.keys())
        raise ValueError(
            f"Unsupported file format '{ext}'. Supported formats: {supported}"
        )

    return extractor(file_path)
